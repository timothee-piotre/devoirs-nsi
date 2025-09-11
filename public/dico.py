from docx import Document

# Load merged words again
with open("dico.txt", "r", encoding="utf-8") as f:
    words1 = f.read().split(",")

with open("dico_fr.txt", "r", encoding="utf-8") as f:
    words2 = f.read().split(",")

# Prepare DOCX document
doc = Document()
doc.add_heading("Dictionnaire fusionné", level=1)

# Ensure we only iterate to the shortest list length
for w1, w2 in zip(words1, words2):
    line = f"{w2.strip()} : {w1.strip()}"  # Word in French first, then the other
    doc.add_paragraph(line)

# Save the document
output_docx = "dico_merged.docx"
doc.save(output_docx)

output_docx
